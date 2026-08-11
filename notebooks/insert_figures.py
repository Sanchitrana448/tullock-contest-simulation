"""
insert_figures.py (v3 - clean rebuild)
---------------------------------------
Rebuilds dissertation_final.docx from dissertation_draft.docx
inserting all 16 figures in correct order in one pass.

Run with: python notebooks/insert_figures.py
"""

import sys
sys.path.insert(0, '.')

import os
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

INPUT_PATH  = "results/dissertation_draft.docx"
OUTPUT_PATH = "results/dissertation_final.docx"
FIGURES_DIR = "results/figures"

# Complete ordered list of all 16 figures
# Each entry: caption text to search for, image filename
# Caption search text must match EXACTLY what is in the draft
FIGURES = [
    {
        "search":   "Figure 1: Synchronous BRD convergence trajectory",
        "filename": "fig01_sync_2player_convergence.png",
    },
    {
        "search":   "Figure 2: Inertial BRD",
        "filename": "fig02_inertial_2player_convergence.png",
    },
    {
        "search":   "Figure 3: Phase portrait",
        "filename": "fig03_phase_portrait_r1.png",
    },
    {
        "search":   "Figure 4: Convergence rate vs. r",
        "filename": "fig04_convergence_rate_vs_r.png",
    },
    {
        "search":   "Figure 5: Mean iterations to convergence vs. r",
        "filename": "fig05_convergence_speed_vs_r.png",
    },
    {
        "search":   "Figure 6: Rent dissipation vs. r",
        "filename": "fig06_rent_dissipation_vs_r.png",
    },
    {
        "search":   "Figure 7: Convergence rate vs. n",
        "filename": "fig07_convergence_rate_vs_n.png",
    },
    {
        "search":   "Figure 8: Mean iterations to convergence vs. n",
        "filename": "fig08_convergence_speed_vs_n.png",
    },
    {
        "search":   "Figure 9: Total effort vs. n",
        "filename": "fig09_total_effort_vs_n.png",
    },
    {
        "search":   "Figure 10: Equilibrium effort vs. Player 2",
        "filename": "fig10_effort_vs_valuation.png",
    },
    {
        "search":   "Figure 11: Equilibrium efforts under symmetric",
        "filename": "fig11_symmetric_vs_asymmetric.png",
    },
    {
        "search":   "Figure 12: Fine-grained convergence rate vs. r",
        "filename": "fig12_fine_r_convergence.png",
    },
    {
        "search":   "Figure 13: Non-convergent trajectory",
        "filename": "fig13_nonconvergent_trajectory.png",
    },
    {
        "search":   "Figure 14: Inertial dynamics at r=3.0",
        "filename": "fig14_inertial_rescue.png",
    },
    {
        "search":   "Figure 15: Full convergence heatmap",
        "filename": "fig15_convergence_heatmap.png",
    },
    {
        "search":   "Figure 16: Convergence rate vs. r for each n",
        "filename": "fig16_convergence_by_n_and_r.png",
    },
]


def find_all_captions(doc):
    """
    Print all italic paragraphs containing Figure
    so we can see exactly what text is in the document.
    """
    print("All figure-related paragraphs found in document:")
    for i, para in enumerate(doc.paragraphs):
        if "Figure" in para.text and len(para.text) > 10:
            print(f"  [{i}] {para.text[:80]}")
    print()


def find_paragraph_index(doc, search_text):
    """Find index of paragraph containing search_text."""
    search_lower = search_text.lower()
    for i, para in enumerate(doc.paragraphs):
        if search_lower in para.text.lower():
            return i
    return None


def insert_image_after(doc, para_index, image_path):
    """Insert image paragraph immediately BEFORE the caption
    at para_index, so the image appears above its caption."""
    from docx.shared import Pt

    caption_para = doc.paragraphs[para_index]

    # Tight spacing: no gap above caption, small gap above image
    caption_para.paragraph_format.space_before = Pt(0)
    caption_para.paragraph_format.space_after  = Pt(0)

    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.space_before = Pt(12)
    img_para.paragraph_format.space_after  = Pt(0)

    run = img_para.add_run()
    run.add_picture(image_path, width=Cm(14))

    # Insert the image paragraph BEFORE the caption
    caption_para._element.addprevious(img_para._element)


print("Opening dissertation draft (clean copy)...")
doc = Document(INPUT_PATH)
print(f"Document loaded: {len(doc.paragraphs)} paragraphs")
print()

# First show all captions found so we can verify
find_all_captions(doc)

# Now insert all figures in reverse order
# (reverse so earlier insertions don't shift later indices)
inserted   = 0
not_found  = []

# Build list of (index, figure) pairs first
to_insert = []
for fig in FIGURES:
    idx = find_paragraph_index(doc, fig["search"])
    if idx is not None:
        to_insert.append((idx, fig))
    else:
        not_found.append(fig["search"])

print(f"Found {len(to_insert)} caption matches out of 16")
if not_found:
    print("NOT FOUND:")
    for s in not_found:
        print(f"  '{s}'")
print()

# Insert in reverse index order
for idx, fig in sorted(to_insert, key=lambda x: x[0],
                        reverse=True):
    image_path = os.path.join(FIGURES_DIR, fig["filename"])
    if os.path.exists(image_path):
        insert_image_after(doc, idx, image_path)
        print(f"Inserted after [{idx}]: {fig['filename']}")
        inserted += 1
    else:
        print(f"IMAGE FILE MISSING: {fig['filename']}")

print()
print(f"Successfully inserted {inserted} figures")

doc.save(OUTPUT_PATH)
print(f"Saved: {OUTPUT_PATH}")
print()
print("=" * 55)
print("Done. Open results/dissertation_final.docx to review.")
print("=" * 55)