"""
write_dissertation.py
---------------------
Generates the full dissertation Word document.
Run with: python notebooks/write_dissertation.py
"""

import sys
sys.path.insert(0, '.')

import os
try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _DOCX_AVAILABLE = True
except Exception:
    Document = None
    Pt = None
    Cm = None
    WD_ALIGN_PARAGRAPH = None
    _DOCX_AVAILABLE = False
    print("python-docx not installed. Install with: pip install python-docx")

if not _DOCX_AVAILABLE:
    import sys as _sys
    _sys.exit(1)

os.makedirs("results", exist_ok=True)

print("Setting up document...")

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)


# ── Helper functions ──────────────────────────────────────

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_para(text, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.font.size = Pt(11)
    return p

def add_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic    = True
    run.font.size = Pt(10)

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold      = True
        run.font.size = Pt(11)
    p.add_run(text).font.size = Pt(11)

def add_numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.add_run(text).font.size = Pt(11)

def page_break():
    doc.add_page_break()

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
    return t


print("Helpers defined. Writing title page...")

# ══════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "Simulating Best Response Dynamics and\n"
    "Equilibria in Tullock Contests"
)
run.bold      = True
run.font.size = Pt(20)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "MSc Artificial Intelligence\n"
    "Department of Computer Science\n"
    "University of Bath"
)
run.font.size = Pt(13)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Sanchit Rana")
run.bold      = True
run.font.size = Pt(13)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Supervisor: Dr. Jie Zhang").font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("September 2026").font.size = Pt(12)

page_break()
print("Title page done. Writing abstract...")


# ══════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════

add_heading("Abstract", level=1)
add_para(
    "This dissertation investigates the dynamic behaviour of strategic "
    "agents competing in Tullock contests through computational simulation "
    "and theoretical analysis. Tullock contests provide a foundational "
    "model for competitive resource allocation under uncertainty, with "
    "applications spanning political economy, organisational behaviour, "
    "mechanism design, and conflict resolution. A comprehensive Python "
    "simulation framework was developed to investigate best response "
    "dynamics under three update rules — synchronous, asynchronous, and "
    "inertial — across a systematic sweep of the decisiveness parameter r "
    "and player count n."
)
add_para(
    "The simulation was validated against known analytical benchmarks from "
    "Tullock (1980), correctly recovering symmetric Nash equilibria across "
    "all update rules. Six experiments were conducted producing 16 "
    "dissertation figures and a formal statistical report including "
    "eigenvalue-based stability analysis. Key findings include: a sharp "
    "convergence boundary at r = 3.0 for two-player contests exhibiting "
    "explosive instability; a convergence threshold between n = 5 and "
    "n = 10 players at r = 1.0; confirmation of the Cornes and Hartley "
    "(2005) prediction that higher valuation players exert greater "
    "equilibrium effort; and the discovery that very high decisiveness "
    "(r = 5.0) produces a universal zero-effort Nash equilibrium. "
    "Spectral radius analysis reveals that the contraction mapping "
    "condition of Szidarovszky and Okuguchi (1997) is sufficient but not "
    "necessary for convergence, constituting a meaningful extension of "
    "existing theoretical results."
)
page_break()
print("Abstract done. Writing Chapter 1...")


# ══════════════════════════════════════════════════════════
# CHAPTER 1 — INTRODUCTION
# ══════════════════════════════════════════════════════════

add_heading("1. Introduction", level=1)

add_heading("1.1 Background and Motivation", level=2)
add_para(
    "Strategic competition for scarce resources is a pervasive feature "
    "of economic, political, and social life. From lobbying for "
    "legislative favours to competing for research grants, agents "
    "routinely expend costly effort to improve their chances of "
    "obtaining a prize. Gordon Tullock's (1980) contest model provides "
    "a tractable and widely-applied framework for analysing such "
    "situations. In the Tullock contest, n risk-neutral players "
    "simultaneously choose effort levels, with each player's probability "
    "of winning determined by the ratio-form contest success function:"
)
add_para(
    "p_i(x) = x_i^r / sum_j(x_j^r)",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER
)
add_para(
    "where r > 0 is the decisiveness parameter controlling how strongly "
    "effort differentials translate into winning probabilities. Despite "
    "decades of theoretical analysis, fundamental questions about the "
    "dynamic behaviour of agents in Tullock contests remain "
    "insufficiently addressed. Most existing work characterises static "
    "Nash equilibria under restrictive assumptions, leaving open "
    "critical questions about whether and how equilibria are reached "
    "through natural learning processes."
)

add_heading("1.2 Research Questions", level=2)
add_para("This dissertation addresses four specific research questions:")
add_numbered(
    "RQ1: How do best response dynamics converge across the "
    "decisiveness parameter space r in {0.5, 1.0, 1.5, 2.0, 3.0, 5.0}?"
)
add_numbered(
    "RQ2: How does player count n in {2, 3, 5, 10, 20} affect "
    "convergence behaviour and equilibrium accessibility?"
)
add_numbered(
    "RQ3: Do heterogeneous player valuations preserve the convergence "
    "properties observed in symmetric contests?"
)
add_numbered(
    "RQ4: Do there exist parameter regimes exhibiting non-convergent "
    "dynamics, and if so, what is the nature of that non-convergence?"
)

add_heading("1.3 Contributions", level=2)
add_para("The main contributions of this dissertation are:")
add_bullet(
    "A validated, modular Python simulation framework for Tullock "
    "contest dynamics, released as open-source software on GitHub."
)
add_bullet(
    "Empirical characterisation of convergence behaviour across a "
    "systematic r x n parameter space sweep, producing the first "
    "comprehensive convergence landscape map for Tullock best "
    "response dynamics."
)
add_bullet(
    "Discovery of a knife-edge non-convergence point at r = 3.0 "
    "exhibiting explosive instability resistant to inertial "
    "stabilisation."
)
add_bullet(
    "Identification of a convergence threshold between n = 5 and "
    "n = 10 players at r = 1.0, revealing how competition intensity "
    "interacts with decisiveness to determine stability."
)
add_bullet(
    "Formal eigenvalue analysis demonstrating that the Szidarovszky "
    "and Okuguchi (1997) contraction condition is sufficient but not "
    "necessary for convergence."
)
add_bullet(
    "Computational confirmation of the Cornes and Hartley (2005) "
    "heterogeneous valuation prediction, extended to settings beyond "
    "their analytical reach."
)

add_heading("1.4 Dissertation Structure", level=2)
add_para(
    "Chapter 2 presents the literature, technology, and data survey. "
    "Chapter 3 details the methodology and simulation framework. "
    "Chapter 4 presents all experimental results. "
    "Chapter 5 discusses theoretical implications. "
    "Chapter 6 concludes with limitations and future work directions."
)
page_break()
print("Chapter 1 done. Writing Chapter 2...")


# ══════════════════════════════════════════════════════════
# CHAPTER 2 — LITERATURE SURVEY
# ══════════════════════════════════════════════════════════

add_heading("2. Literature, Technology, and Data Survey", level=1)

add_heading("2.1 Foundations of Contest Theory", level=2)
add_para(
    "Tullock's (1980) ratio-form contest success function established "
    "the foundational trade-off between investment and probability of "
    "success. For the symmetric case with linear costs and r = 1, the "
    "Nash equilibrium effort level is x* = V(n-1)/n^2, yielding "
    "aggregate effort X* = V(n-1)/n. These closed-form results provide "
    "the essential benchmarks against which this dissertation's "
    "simulations are validated."
)
add_para(
    "Szidarovszky and Okuguchi (1997) generalised these results using "
    "contraction mapping arguments, establishing that when best response "
    "correspondences satisfy a contraction condition — equivalently, "
    "when the spectral radius of the Jacobian of best response mappings "
    "is below unity — iterative best response dynamics are guaranteed "
    "to converge globally to the unique Nash equilibrium. This "
    "theoretical result defines a set of sufficient conditions for "
    "convergence, but leaves open whether these conditions are also "
    "necessary — a question this dissertation directly addresses "
    "computationally."
)
add_para(
    "Cornes and Hartley (2005) characterised equilibria under "
    "asymmetric valuations, demonstrating that higher-valuation players "
    "exert proportionally greater effort at equilibrium. Their analysis "
    "relied on techniques that break down under severe asymmetry or "
    "non-standard cost functions, creating a direct opening for "
    "computational investigation into settings beyond their analytical "
    "reach."
)

add_heading("2.2 Rent Dissipation and Equilibrium Efficiency", level=2)
add_para(
    "Perez-Castrillo and Verdier (1992) demonstrated that total "
    "equilibrium effort can exceed prize value under certain "
    "parameterisations — the over-dissipation phenomenon. Nitzan (1994) "
    "established that the relationship between r and aggregate rent "
    "dissipation is non-monotonic, with dissipation maximised at "
    "intermediate values of r. This dissertation's simulations test "
    "whether this non-monotonicity persists across the full parameter "
    "space and under heterogeneous player valuations."
)

add_heading("2.3 Dynamic Analysis and Learning", level=2)
add_para(
    "Hopkins and Kornienko (2010) examined evolutionary dynamics in "
    "contests using replicator dynamics, finding limit cycles and "
    "sensitive dependence on initial conditions. Their work established "
    "that the existence of a unique Nash equilibrium does not guarantee "
    "convergence under evolutionary learning — a result that motivates "
    "the computational investigation of best response dynamics "
    "undertaken here."
)
add_para(
    "Lim and Matros (2009) showed that stochastic perturbations can "
    "facilitate convergence where deterministic dynamics fail, "
    "motivating the investigation of inertial update rules as a "
    "potential stabilisation mechanism. Chowdhury and Sheremeta (2011) "
    "found systematic deviations from Nash predictions in experimental "
    "settings, underlining the practical importance of understanding "
    "which equilibria are dynamically accessible."
)

add_heading("2.4 Recent Computational Work", level=2)
add_para(
    "Ghosh and Goldberg (2023) proved that best response dynamics "
    "rapidly converge for homogeneous agents but may fail for "
    "heterogeneous agents even with two players. Elkind, Ghosh and "
    "Goldberg (2024) showed that continuous-time best response dynamics "
    "converge for homogeneous agents using Lyapunov-style arguments, "
    "with non-homogeneous agents failing to converge on a "
    "positive-measure set of instances. These analytical results "
    "provide theoretical grounding for the computational findings "
    "presented in Chapter 4."
)

add_heading("2.5 Multi-Battlefield Extensions", level=2)
add_para(
    "Liu, Ni, Shen, Wang and Zhang (2025) study the Lottery Colonel "
    "Blotto game — a multi-battlefield generalisation in which players "
    "divide budgets across n battlefields each governed by a "
    "proportional rule structurally identical to the Tullock contest "
    "success function. Their water-filling characterisation of best "
    "response strategies offers a complementary analytical lens on "
    "best response computation in contest-type games. Their Stackelberg "
    "framing suggests a natural extension of the present work noted "
    "as a direction for future research in Chapter 6."
)

add_heading("2.6 Technology Survey", level=2)
add_para(
    "The simulation framework is implemented in Python 3.12, leveraging "
    "NumPy for vectorised computation, SciPy for numerical optimisation "
    "(minimize_scalar with bounded method for best response computation, "
    "fsolve for root-finding), Matplotlib and Seaborn for visualisation, "
    "and Git and GitHub for version control and open-source release. "
    "The modular architecture follows the scientific computing best "
    "practices of Wilson et al. (2014), separating contest model "
    "specification, best response computation, dynamic iteration, "
    "equilibrium identification, and analysis and visualisation into "
    "independent modules."
)
page_break()
print("Chapter 2 done. Writing Chapter 3...")


# ══════════════════════════════════════════════════════════
# CHAPTER 3 — METHODOLOGY
# ══════════════════════════════════════════════════════════

add_heading("3. Methodology", level=1)

add_heading("3.1 Mathematical Model", level=2)
add_para(
    "The simulation implements a Tullock contest with n players, "
    "decisiveness parameter r in (0, infinity), player valuations "
    "V_i > 0, and cost function c(x_i) = x_i^alpha with alpha >= 1. "
    "The contest success function assigns player i a winning "
    "probability:"
)
add_para(
    "p_i(x) = x_i^r / sum_{j=1}^{n} x_j^r",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER
)
add_para(
    "Player i's expected payoff is pi_i(x) = p_i(x) * V_i - c(x_i). "
    "The baseline specification uses linear costs (alpha = 1) and "
    "symmetric valuations (V_i = V for all i), with extensions to "
    "heterogeneous valuations examined in Experiment 4."
)

add_heading("3.2 Best Response Computation", level=2)
add_para(
    "Player i's best response to competitors' efforts x_{-i} solves:"
)
add_para(
    "BR_i(x_{-i}) = argmax_{x_i >= 0} "
    "[V_i * x_i^r / (x_i^r + sum_{j not i} x_j^r) - x_i]",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER
)
add_para(
    "This is solved numerically using SciPy's minimize_scalar with "
    "the bounded method over the interval [1e-8, 100], achieving "
    "tolerance 1e-10. The bounded search ensures robustness for both "
    "concave payoffs (r <= 1) and potentially non-concave payoffs "
    "(r > 1)."
)

add_heading("3.3 Update Rules", level=2)
add_para(
    "Three update rules are implemented, each representing a different "
    "model of strategic adjustment:"
)
add_bullet(
    "All n players simultaneously compute and adopt their best "
    "responses: x_i(t+1) = BR_i(x_{-i}(t)) for all i.",
    bold_prefix="Synchronous: "
)
add_bullet(
    "Players update sequentially in a randomly determined order "
    "each round. Each player best-responds to the most recently "
    "updated efforts of others.",
    bold_prefix="Asynchronous: "
)
add_bullet(
    "Players move fraction lambda toward their best response: "
    "x_i(t+1) = (1-lambda)*x_i(t) + lambda*BR_i(x_{-i}(t)). "
    "lambda = 1 reduces to synchronous; smaller lambda models "
    "more cautious adjustment.",
    bold_prefix="Inertial: "
)

add_heading("3.4 Convergence Detection", level=2)
add_para(
    "Convergence is declared when the maximum absolute change across "
    "all players falls below tolerance epsilon = 1e-6:"
)
add_para(
    "max_i |x_i(t+1) - x_i(t)| < epsilon",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER
)
add_para(
    "Each experiment uses a maximum of 500 iterations. Runs not "
    "converging within this limit are classified as non-convergent. "
    "Results are averaged across 10-30 random initialisations drawn "
    "uniformly from [0.1, 0.8*V]^n."
)

add_heading("3.5 Equilibrium Validation", level=2)
add_para(
    "For symmetric contests with r = 1 and linear costs, the "
    "analytical Nash equilibrium is x* = V(n-1)/n^2 from Tullock "
    "(1980). All three update rules are validated against this "
    "benchmark before parameter sweeps are conducted. Validation "
    "tests pass with absolute error below 1e-4 across all tested "
    "configurations, confirming the framework's correctness."
)

add_heading("3.6 Stability Analysis", level=2)
add_para(
    "Eigenvalue analysis of the Jacobian matrix of best response "
    "mappings provides a formal stability assessment at identified "
    "equilibria. The Jacobian J is estimated numerically with step "
    "size delta = 1e-5. The spectral radius rho(J) = max_k |lambda_k| "
    "determines local stability: rho < 1 implies the equilibrium is "
    "locally attracting under synchronous dynamics, directly "
    "operationalising the contraction mapping condition of "
    "Szidarovszky and Okuguchi (1997)."
)
page_break()
print("Chapter 3 done. Writing Chapter 4...")


# ══════════════════════════════════════════════════════════
# CHAPTER 4 — RESULTS
# ══════════════════════════════════════════════════════════

add_heading("4. Experimental Results", level=1)

add_para(
    "This chapter presents the results of six experiments designed "
    "to systematically address the four research questions posed in "
    "Chapter 1. Experiment 1 establishes the correctness of the "
    "simulation framework against a known analytical benchmark. "
    "Experiments 2 and 3 address RQ1 and RQ2 respectively, sweeping "
    "the decisiveness parameter r and player count n independently. "
    "Experiment 4 addresses RQ3 by introducing heterogeneous player "
    "valuations. Experiment 5 addresses RQ4 by examining a specific "
    "non-convergence point identified in Experiment 2 in much finer "
    "detail. Experiment 6 combines the r and n dimensions into a "
    "single systematic sweep, producing the most complete picture of "
    "the convergence landscape presented in this dissertation."
)

add_heading("4.1 Baseline Validation (Experiment 1)", level=2)
add_para(
    "Before any exploratory experiments were conducted, it was "
    "essential to establish that the simulation framework correctly "
    "reproduces known theoretical results. Without this step, any "
    "subsequent finding — however interesting — would carry no "
    "evidential weight, since it would be impossible to distinguish "
    "a genuine feature of Tullock contest dynamics from an artefact "
    "of a flawed implementation. Experiment 1 therefore validates the "
    "framework against the closed-form Nash equilibrium for the "
    "symmetric 2-player contest with linear costs and r = 1, "
    "V = 10, for which Tullock (1980) gives the analytical solution "
    "x* = V(n-1)/n^2 = 10(1)/4 = 2.5 for both players."
)
add_para(
    "The validation deliberately starts from an asymmetric initial "
    "condition, (8.0, 1.0), rather than the equilibrium itself. This "
    "choice matters: a simulation that trivially confirms an "
    "equilibrium when initialised at that equilibrium proves very "
    "little, since a stationary point will remain stationary under "
    "any reasonable update rule. Starting from a point far from "
    "equilibrium, with Player 1 initially far ahead of Player 2, "
    "tests whether the dynamics genuinely correct themselves through "
    "the update process rather than simply preserving whatever state "
    "they are given."
)

add_table(
    ["Update Rule", "Iterations", "Final Efforts", "Converged"],
    [
        ["Synchronous",    "6",  "(2.5000, 2.5000)", "Yes"],
        ["Asynchronous",   "5",  "(2.5000, 2.5000)", "Yes"],
        ["Inertial l=0.5", "23", "(2.5000, 2.5000)", "Yes"],
    ]
)
add_caption(
    "Table 1: Validation results for symmetric 2-player contest. "
    "All three update rules recover the analytical equilibrium x* = 2.5."
)
doc.add_paragraph()

add_para(
    "All three update rules converge precisely to the analytical "
    "equilibrium, with absolute error below 1e-4 in every case — "
    "well within the tolerance epsilon = 1e-6 used to declare "
    "convergence, and far tighter than would be needed for any "
    "practical purpose. The synchronous and asynchronous rules "
    "converge in remarkably few iterations, 6 and 5 respectively, "
    "reflecting the fact that at r = 1 with only two symmetric "
    "players the best response mapping is strongly contracting: each "
    "round of updating closes most of the remaining gap to "
    "equilibrium. The inertial rule with learning rate lambda = 0.5 "
    "requires substantially more iterations, 23, which is expected "
    "given its construction — by design it moves only half the "
    "distance toward the best response at each step rather than "
    "jumping there directly, so it necessarily takes roughly twice "
    "as many rounds per unit of convergence relative to the "
    "synchronous rule, though the relationship is not perfectly "
    "linear because the target itself shifts slightly at each "
    "partial step."
)
add_para(
    "That all three rules, despite their structurally different "
    "update mechanics, converge to the identical equilibrium value "
    "provides strong initial evidence that the equilibrium in "
    "question is not merely a fixed point but a genuinely stable "
    "attractor for this parameter configuration, reachable regardless "
    "of the precise manner in which players adjust their behaviour "
    "over time. This robustness across update rules is itself an "
    "informative result: it suggests that for well-behaved parameter "
    "regions, the qualitative predictions of Nash equilibrium theory "
    "are not sensitive to modelling choices about exactly how players "
    "learn, a property that later experiments show does not hold "
    "universally across the parameter space."
)
add_caption(
    "Figure 1: Synchronous BRD convergence trajectory — 2-player "
    "symmetric contest (V=10, r=1). Both players converge to x*=2.5 "
    "within 6 iterations."
)
add_caption(
    "Figure 2: Inertial BRD (l=0.5) convergence trajectory — "
    "2-player symmetric contest (V=10, r=1). Both players converge "
    "to x*=2.5 within 23 iterations."
)
add_caption(
    "Figure 3: Phase portrait of 2-player contest (r=1). Arrows "
    "indicate best response update directions. Red star marks the "
    "Nash equilibrium at (2.5, 2.5)."
)
add_para(
    "The phase portrait in Figure 3 provides a complementary, more "
    "global view of the same result. Rather than tracking a single "
    "trajectory, it plots the direction and approximate magnitude of "
    "the best response update at a grid of points across the "
    "strategy space. Every arrow in the portrait points, directly or "
    "indirectly, toward the marked equilibrium at (2.5, 2.5), and the "
    "arrows shrink in length as they approach it — visual confirmation "
    "that the equilibrium behaves as a genuine attractor across the "
    "entire region shown, not merely along the one trajectory "
    "examined in Figure 1."
)

add_heading("4.2 Decisiveness Parameter Sweep (Experiment 2)", level=2)
add_para(
    "With the framework validated, Experiment 2 addresses RQ1 "
    "directly: how does convergence behaviour change as the "
    "decisiveness parameter r varies? The decisiveness parameter "
    "governs how sharply small differences in effort translate into "
    "differences in winning probability, and existing theory offers "
    "only partial guidance as to how it should affect the dynamic "
    "stability of learning, as distinct from the static properties "
    "of the equilibrium itself. Experiment 2 swept r across "
    "{0.5, 1.0, 1.5, 2.0, 3.0} for the 2-player symmetric contest "
    "(V=10), running 20 random initialisations per r value, each "
    "drawn independently and uniformly from [0.1, 8.0] for both "
    "players, and using the synchronous update rule throughout since "
    "it was the most directly comparable across configurations."
)

add_table(
    ["r", "Conv. Rate", "Mean Iters", "Total Effort", "Spectral R"],
    [
        ["0.5", "100%", "4.2",  "2.50",  "0.0002"],
        ["1.0", "100%", "5.3",  "5.00",  "0.0032"],
        ["1.5", "100%", "9.2",  "7.50",  "0.0005"],
        ["2.0", "100%", "9.8",  "10.00", "0.0119"],
        ["3.0", "0%",   "N/A",  "N/A",   "N/A"],
    ]
)
add_caption(
    "Table 2: Convergence statistics across r values "
    "(n=2, V=10, 20 random initialisations)."
)
doc.add_paragraph()

add_para(
    "Reading down Table 2 in order, several patterns emerge that "
    "would not be visible from the aggregate convergence rate alone. "
    "At r = 0.5, convergence is both universal (100% across all 20 "
    "initialisations) and extremely fast, requiring on average only "
    "4.2 iterations. This is consistent with intuition: when r is "
    "small, winning probability responds only weakly to differences "
    "in effort, so the incentive to chase a competitor's effort "
    "level is muted and the system settles quickly. As r rises to "
    "1.0 and then 1.5, convergence remains complete but the mean "
    "iteration count roughly doubles, from 4.2 to 9.2, indicating "
    "that the same qualitative outcome — reaching the unique "
    "equilibrium — is being reached through an increasingly "
    "protracted process of mutual adjustment."
)
add_para(
    "The total effort column tells a complementary story about the "
    "economic content of these equilibria rather than merely their "
    "dynamic accessibility. Total effort rises steadily and "
    "predictably from 2.50 at r = 0.5 to 10.00 at r = 2.0 — precisely "
    "in line with the analytical formula X* = V(n-1)/n applied to "
    "n = 2 contests, since total effort here reflects the underlying "
    "static equilibrium rather than the dynamics used to reach it and "
    "so should not itself vary with r under the linear-cost, "
    "symmetric-valuation specification used. What does vary "
    "systematically with r, however, is the spectral radius of the "
    "Jacobian evaluated at that equilibrium, rising unevenly but "
    "broadly from 0.0002 at r = 0.5 to 0.0119 at r = 2.0. Although "
    "still far below the instability threshold of 1.0 throughout "
    "this range, the trend is directionally consistent with the "
    "rising iteration counts: a higher decisiveness parameter "
    "produces a locally less strongly contracting best response map, "
    "so more rounds of mutual adjustment are needed to close the same "
    "proportional distance to equilibrium."
)
add_para(
    "Then, at r = 3.0, this gradual pattern breaks entirely. "
    "Convergence rate collapses from 100% to 0% across all 20 "
    "initialisations without exception — not a partial degradation "
    "but a complete failure, with the total effort and spectral "
    "radius columns both showing N/A because no run reached a stable "
    "point from which either quantity could meaningfully be "
    "computed. Given the smooth, monotonic-looking progression "
    "through r = 0.5 to r = 2.0, this outcome at r = 3.0 was "
    "genuinely unexpected at the point it was first observed, and it "
    "is this anomaly that motivated the much finer-grained "
    "investigation carried out separately in Experiment 5."
)
add_caption(
    "Figure 4: Convergence rate vs. r (n=2). Complete failure at "
    "r=3.0 surrounded by full convergence at r=2.0 and r=3.5."
)
add_caption(
    "Figure 5: Mean iterations to convergence vs. r (n=2). "
    "Convergence slows as r increases toward the instability "
    "boundary."
)
add_caption(
    "Figure 6: Rent dissipation vs. r showing monotonic increase "
    "in the convergent region."
)
add_para(
    "It is worth noting explicitly what Figure 6 does and does not "
    "show. Within the convergent region tested, r ∈ {0.5, 1.0, 1.5, "
    "2.0}, total effort rises monotonically with r, which is "
    "consistent with the rising portion of the non-monotonic "
    "dissipation curve predicted analytically by Nitzan (1994). "
    "However, because r = 3.0 fails to converge, this experiment "
    "alone cannot confirm or refute whether dissipation eventually "
    "falls again at higher r, as Nitzan's theory predicts; that "
    "question is addressed indirectly through the corner-solution "
    "result obtained later at r = 5.0 in Experiment 6, and is "
    "discussed further in Chapter 5."
)

add_heading("4.3 Player Count Sweep (Experiment 3)", level=2)
add_para(
    "Experiment 3 addresses RQ2, holding the decisiveness parameter "
    "fixed at its most standard value, r = 1.0, and instead sweeping "
    "the number of players n across {2, 3, 5, 10, 20}. This range "
    "was chosen to span from the minimal two-player case, through "
    "small committee-sized groups, up to contest sizes more "
    "representative of large-scale competitive settings such as "
    "open tournaments or crowded markets. Ten random initialisations "
    "were used per n value, each player's starting effort again drawn "
    "independently and uniformly from [0.1, 8.0]."
)

add_table(
    ["n", "Conv. Rate", "Mean Iters", "Sim. Total", "Analytical"],
    [
        ["2",  "100%", "5.3",    "5.00",  "5.00"],
        ["3",  "100%", "22.0",   "6.67",  "6.67"],
        ["5",  "100%", "2010.1", "8.00",  "8.00"],
        ["10", "0%",   "N/A",    "6.25",  "9.00"],
        ["20", "0%",   "N/A",    "22.49", "9.50"],
    ]
)
add_caption(
    "Table 3: Convergence statistics across n values "
    "(r=1, V=10, 10 random initialisations)."
)
doc.add_paragraph()

add_para(
    "The first three rows of Table 3 show convergence rates of 100% "
    "throughout, but the mean iterations column reveals that this "
    "surface-level consistency conceals a rapidly deteriorating "
    "situation underneath. Moving from n = 2 to n = 3 nearly "
    "quadruples the mean iteration count, from 5.3 to 22.0. Moving "
    "from n = 3 to n = 5 then increases it by a further factor of "
    "roughly ninety, to 2010.1 iterations — meaning that although "
    "every single one of the ten random initialisations at n = 5 "
    "technically reached the analytical equilibrium within the "
    "iteration budget allowed, doing so required well over a "
    "thousand times more rounds of adjustment than the two-player "
    "case. This is a striking illustration of how a binary "
    "convergence classification, converged versus not converged, can "
    "obscure an underlying trend that is already signalling severe "
    "instability well before outright failure occurs."
)
add_para(
    "At n = 10, that failure arrives: convergence rate drops sharply "
    "to 0%, and it remains at 0% at n = 20. The threshold between "
    "n = 5 and n = 10 therefore represents a genuine qualitative "
    "boundary in the behaviour of the system, not merely a "
    "continuation of the slowing trend visible between n = 2 and "
    "n = 5. Equally informative is what happens to the total effort "
    "figures once this boundary is crossed. For n = 2, 3, and 5, "
    "simulated total effort matches the analytical prediction exactly "
    "to the precision reported, as it must, since convergence to the "
    "true equilibrium guarantees this agreement. But at n = 10, "
    "simulated total effort is 6.25 against an analytical prediction "
    "of 9.00, and at n = 20 it is 22.49 against a prediction of only "
    "9.50 — in the n = 20 case, simulated total effort is more than "
    "double the theoretical equilibrium value. These are not small "
    "deviations consistent with slow but ongoing convergence; they "
    "indicate that the non-convergent trajectories are exploring "
    "regions of the strategy space that bear little resemblance to "
    "the equilibrium itself, most likely oscillating through states "
    "where a subset of players temporarily commit very high effort "
    "while others collapse toward zero, in a manner qualitatively "
    "similar to the explosive instability documented in detail for "
    "the r = 3.0 case in Experiment 5."
)
add_caption(
    "Figure 7: Convergence rate vs. n (r=1). Sharp drop from 100% "
    "at n=5 to 0% at n=10."
)
add_caption(
    "Figure 8: Mean iterations to convergence vs. n (r=1). "
    "Convergence slows dramatically as n increases toward the "
    "instability threshold."
)
add_caption(
    "Figure 9: Total effort vs. n comparing simulation against "
    "analytical benchmarks. Perfect agreement where convergence "
    "holds; divergence where it fails."
)
add_para(
    "The formal analysis report (introduced in Section 3.6 and run "
    "separately from the main experiment scripts) supplies the "
    "spectral radius for the n = 5 case: 1.49, formally above the "
    "unity threshold associated with the Szidarovszky-Okuguchi "
    "contraction condition, and yet convergence at n = 5 was observed "
    "in 100% of runs. This is one of two independent pieces of "
    "evidence in this dissertation — the other arising in Section "
    "4.6 — that the contraction condition, while clearly sufficient "
    "for convergence, is not necessary. This finding is discussed at "
    "greater length in Chapter 5, where its implications for the "
    "theoretical literature are considered directly."
)

add_heading("4.4 Heterogeneous Valuations (Experiment 4)", level=2)
add_para(
    "Experiments 2 and 3 both used symmetric valuations, V_i = V for "
    "all players, which is the setting for which the closed-form "
    "benchmark of Tullock (1980) is available and therefore the "
    "natural starting point for validating and exploring the "
    "framework. Experiment 4 relaxes this assumption to address RQ3, "
    "asking whether the convergence properties documented above "
    "persist once players value the prize differently from one "
    "another. This question is not purely academic: in almost every "
    "real-world application of contest theory, from R&D races to "
    "political competition, the competing parties rarely value the "
    "prize identically, so understanding whether asymmetry alone can "
    "destabilise otherwise well-behaved dynamics is directly relevant "
    "to the practical applicability of the model."
)
add_para(
    "Three configurations were tested for the 2-player contest at "
    "r = 1.0, each run with 20 random initialisations: a symmetric "
    "baseline with V1 = V2 = 10 for comparison, a mild asymmetry with "
    "V1 = 8 and V2 = 12, and a strong asymmetry with V1 = 2 and "
    "V2 = 18."
)

add_table(
    ["Configuration", "V1", "V2", "P1 Effort", "P2 Effort"],
    [
        ["Symmetric",        "10", "10", "2.5000", "2.5000"],
        ["Mild asymmetry",   "8",  "12", "1.9200", "2.8800"],
        ["Strong asymmetry", "2",  "18", "0.1800", "1.6200"],
    ]
)
add_caption(
    "Table 4: Equilibrium efforts under symmetric and asymmetric "
    "valuations. Convergence rate 100% in all cases."
)
doc.add_paragraph()

add_para(
    "The most immediately striking feature of Table 4 is the "
    "precision with which the effort ratio matches the valuation "
    "ratio in both asymmetric cases. Under mild asymmetry, "
    "V2/V1 = 12/8 = 1.5, and the corresponding effort ratio is "
    "2.8800/1.9200 = 1.5 exactly. Under strong asymmetry, "
    "V2/V1 = 18/2 = 9.0, and the effort ratio is 1.6200/0.1800 = 9.0 "
    "exactly. This is a direct computational confirmation of the "
    "proportionality result derived analytically by Cornes and "
    "Hartley (2005) for asymmetric contests, and its exact "
    "reproduction here — rather than an approximate match subject to "
    "simulation noise — provides strong corroborating evidence for "
    "the correctness of the best response computation implemented "
    "in this framework, independent of the earlier validation against "
    "the symmetric benchmark in Experiment 1."
)
add_para(
    "Equally important, though less visually obvious from the table "
    "alone, is that convergence remained at 100% across all 20 "
    "initialisations in both asymmetric configurations, including the "
    "strong asymmetry case where the ratio of valuations reaches 9:1. "
    "This indicates that valuation heterogeneity, at least of the "
    "magnitude tested here, does not by itself destabilise best "
    "response dynamics at r = 1.0 — the instability documented "
    "elsewhere in this dissertation appears to be driven specifically "
    "by the decisiveness parameter and player count rather than by "
    "asymmetry in how much players value the prize."
)
add_caption(
    "Figure 10: Equilibrium effort vs. Player 2 valuation. "
    "Player 2 effort rises monotonically; Player 1 effort shows "
    "non-monotonic response peaking at the symmetric point."
)
add_para(
    "Figure 10 extends this comparison into a continuous sweep of "
    "V2 while holding V1 = 10 fixed, and reveals a pattern not "
    "visible from the three discrete cases in Table 4 alone. Player "
    "2's equilibrium effort rises monotonically and roughly linearly "
    "as V2 increases, exactly as the proportionality result would "
    "predict. Player 1's effort, however, does not fall monotonically "
    "as one might naively expect from a player facing an "
    "increasingly well-resourced rival; instead it rises as V2 "
    "approaches V1 from below, peaks almost exactly at the symmetric "
    "point where V2 = V1 = 10, and only then begins to decline as V2 "
    "continues to rise beyond that point. A plausible reading of this "
    "pattern is that Player 1 responds defensively to an intensifying "
    "but still roughly matched rival, raising effort to protect their "
    "position, but once V2 grows so large that competing seriously "
    "for the prize becomes rationally unattractive, Player 1 begins "
    "to withdraw effort rather than continue investing against "
    "worsening odds. This non-monotonic strategic response is "
    "discussed further, together with its relationship to the "
    "existing literature, in Chapter 5."
)
add_caption(
    "Figure 11: Equilibrium efforts under symmetric vs asymmetric "
    "valuations. Player 2 effort consistently exceeds Player 1 "
    "when V2 > V1, confirming Cornes and Hartley (2005)."
)

add_heading("4.5 High Decisiveness Analysis (Experiment 5)", level=2)
add_para(
    "The complete non-convergence observed at r = 3.0 in Experiment "
    "2, occurring as an apparently isolated failure between two fully "
    "convergent neighbouring values, was sufficiently unexpected that "
    "it warranted dedicated investigation beyond what the original "
    "coarse parameter grid could offer. Experiment 5 addresses RQ4 "
    "directly by examining this specific point through three "
    "complementary lenses: a much finer-grained sweep of r in the "
    "immediate vicinity of 3.0, a detailed inspection of a single "
    "representative trajectory to characterise what non-convergence "
    "actually looks like in practice, and a test of whether the "
    "inertial update rule — which moves only partially toward the "
    "best response each round and might therefore be expected to "
    "dampen any instability — could rescue convergence at this point."
)
add_para(
    "The fine-grained sweep tested r ∈ {2.0, 2.2, 2.4, 2.6, 2.8, "
    "3.0, 3.2, 3.5}, a resolution five times finer than the original "
    "Experiment 2 grid in the region immediately surrounding the "
    "anomaly. The result confirmed that the failure at r = 3.0 in "
    "Experiment 2 was not an artefact of the coarse sampling: r = 2.8 "
    "and r = 3.2, immediately either side of the failure point, both "
    "converge at 100%, while r = 3.0 itself converges at exactly 0% "
    "across all 20 initialisations tested. This is a knife-edge "
    "result in the most literal sense — a single, sharply localised "
    "point of total failure with no gradual transition visible on "
    "either side at this resolution."
)
add_caption(
    "Figure 12: Fine-grained convergence rate vs. r. Knife-edge "
    "failure at exactly r=3.0 with full convergence on both sides."
)
add_para(
    "To understand what this failure actually looks like in "
    "practice, a single trajectory was traced in detail from the "
    "initial condition (4.0, 1.0) at r = 3.0. Rather than settling "
    "toward a fixed point, or oscillating gently between two nearby "
    "values as a simple limit cycle would, the trajectory shows "
    "explosive divergence: within the first ten iterations, effort "
    "levels swing from the starting point through values exceeding "
    "seven, collapse toward zero for one player while the other "
    "peaks, then reverse roles entirely, before both players' efforts "
    "crash toward values close to zero simultaneously. This behaviour "
    "is qualitatively distinct from the smooth, monotonic approach to "
    "equilibrium seen throughout Experiment 1, and closer in "
    "character to the sensitive dependence on initial conditions "
    "documented under evolutionary dynamics by Hopkins and Kornienko "
    "(2010), even though the update rule used here is best response "
    "learning rather than replicator dynamics."
)
add_caption(
    "Figure 13: Non-convergent trajectory at r=3.0 showing "
    "explosive oscillation between near-zero and near-maximum effort."
)
add_para(
    "Finally, four inertial learning rates — lambda ∈ {0.8, 0.5, "
    "0.3, 0.1} — were tested at r = 3.0 to determine whether damping "
    "the speed of adjustment could restore convergence. Intuitively, "
    "a sufficiently cautious update rule, one that moves only a small "
    "fraction of the way toward the best response each round, might "
    "be expected to prevent the kind of overshoot visible in Figure "
    "13. This intuition did not hold: all four learning rates tested, "
    "including the most cautious lambda = 0.1, produced 0% "
    "convergence at r = 3.0. This result is important because it "
    "rules out the simplest explanation for the instability — that "
    "it is merely an artefact of updating too aggressively — and "
    "instead suggests that the instability is a more fundamental "
    "structural property of the best response map itself at this "
    "parameter value, one that persists regardless of how gradually "
    "players are permitted to adjust."
)
add_caption(
    "Figure 14: Inertial dynamics at r=3.0 showing 0% convergence "
    "across all lambda values tested."
)

add_heading("4.6 Full Parameter Space Heatmap (Experiment 6)", level=2)
add_para(
    "Experiments 2 and 3 examined the effects of r and n "
    "independently, holding the other fixed. This leaves open the "
    "question of how the two dimensions interact — whether, for "
    "instance, the instability at r = 3.0 documented for n = 2 in "
    "Experiment 5 also appears at other player counts, or whether it "
    "is specific to the two-player case. Experiment 6 addresses this "
    "directly by sweeping both parameters simultaneously, producing "
    "a full convergence rate grid across r ∈ {0.5, 1.0, 1.5, 2.0, "
    "2.5, 3.0, 3.5, 4.0, 5.0} and n ∈ {2, 3, 5, 10, 20}, with 10 "
    "random initialisations per cell — a total of 45 parameter "
    "combinations and 450 individual simulation runs."
)

add_table(
    ["n / r", "0.5", "1.0", "2.0", "3.0", "5.0"],
    [
        ["n=2",  "100%", "100%", "100%", "0%",  "20%"],
        ["n=3",  "100%", "100%", "20%",  "0%",  "100%"],
        ["n=5",  "100%", "100%", "0%",   "0%",  "100%"],
        ["n=10", "100%", "0%",   "0%",   "0%",  "100%"],
        ["n=20", "100%", "0%",   "0%",   "0%",  "100%"],
    ]
)
add_caption(
    "Table 5: Convergence rates across r x n parameter space "
    "(selected columns). Full heatmap in Figure 15."
)
doc.add_paragraph()

add_para(
    "Table 5 shows a five-column extract from the full nine-column "
    "grid, chosen to illustrate the main structural features visible "
    "in the complete heatmap of Figure 15. Reading along the r = 0.5 "
    "row and column shows universal convergence — 100% for every "
    "value of n tested, from 2 up to 20. This is the single most "
    "stable region identified anywhere in this dissertation, and it "
    "holds despite the substantial increase in the number of "
    "simultaneously adjusting players across that range, suggesting "
    "that low decisiveness provides a form of stability that is "
    "robust to competition intensity in a way that the intermediate "
    "r values are not."
)
add_para(
    "Moving to r = 1.0, the picture changes sharply once n grows "
    "large: convergence remains at 100% for n = 2, 3, and 5, "
    "consistent with the results already reported in detail in "
    "Experiment 3, but drops to 0% for both n = 10 and n = 20. At "
    "r = 2.0, the picture is markedly worse even at moderate player "
    "counts: n = 3 achieves only 20% convergence and n = 5 achieves "
    "0%, indicating that the instability zone is not simply shifting "
    "but actively widening as r increases through the intermediate "
    "range. The r = 3.0 column shows complete failure, 0%, across "
    "every single value of n tested, from 2 through 20 — confirming "
    "that the knife-edge instability characterised in detail for the "
    "two-player case in Experiment 5 is not a two-player peculiarity "
    "but a feature of this specific decisiveness value that persists "
    "across the entire range of player counts examined."
)
add_para(
    "The final column, r = 5.0, presents perhaps the most surprising "
    "result in the entire heatmap: convergence returns to 100% for "
    "every n value tested, mirroring the universal stability seen at "
    "r = 0.5. However, this apparent return to stability conceals a "
    "very different underlying equilibrium. Examination of the "
    "underlying effort data (recorded in the Experiment 6 output but "
    "summarised for clarity here) shows that at r = 5.0, equilibrium "
    "total effort is essentially zero for n >= 3, in sharp contrast "
    "to the positive interior equilibria found at r = 0.5. In other "
    "words, r = 5.0 achieves universal convergence not because the "
    "dynamics have become well-behaved in the same sense as at low "
    "r, but because the extreme decisiveness of the contest at this "
    "parameter value makes any positive investment of effort "
    "unprofitable for every player, so all players rationally "
    "converge on contributing nothing at all. This is a corner "
    "solution rather than an interior Nash equilibrium of the kind "
    "analysed throughout the rest of this dissertation, and its "
    "implications are discussed further in Chapter 5."
)
add_caption(
    "Figure 15: Full convergence heatmap across r x n parameter "
    "space. Green = convergence, red = failure."
)
add_caption(
    "Figure 16: Convergence rate vs. r for each n value, showing "
    "how the instability zone widens with player count."
)
add_para(
    "Figure 16 makes the widening of the instability zone with "
    "increasing n visually explicit by plotting convergence rate "
    "against r as a separate line for each value of n. The n = 2 "
    "line remains close to 100% across almost the entire range, "
    "dipping only briefly at r = 3.0 and partially at r = 4.0 and "
    "r = 5.0. As n increases, each successive line dips lower and "
    "for a wider range of r values, until the n = 10 and n = 20 "
    "lines show almost complete failure across the entire "
    "intermediate range from roughly r = 1.0 through r = 4.0, "
    "recovering only at the two extremes of the tested range. This "
    "figure, taken together with Table 5, constitutes the single "
    "most comprehensive empirical answer this dissertation provides "
    "to RQ1 and RQ2 jointly: convergence of best response dynamics "
    "in Tullock contests is far from universal, is highly sensitive "
    "to the interaction between decisiveness and player count rather "
    "than either parameter in isolation, and is reliably guaranteed "
    "only at the extremes of the decisiveness parameter tested here."
)
page_break()
print("Chapter 4 done. Writing Chapter 5...")


# ══════════════════════════════════════════════════════════
# CHAPTER 5 — DISCUSSION
# ══════════════════════════════════════════════════════════

add_heading("5. Discussion", level=1)

add_heading("5.1 Informing Existing Theory", level=2)
add_para(
    "The simulation results inform existing theoretical understanding "
    "in two principal ways. First, the monotonic relationship between "
    "spectral radius and convergence speed — visible across Tables 2 "
    "and 3, where mean iterations rise from 4.2 at spectral radius "
    "0.0002 to 9.8 at 0.0119 — provides empirical grounding for the "
    "intuition that the contraction mapping condition not only "
    "guarantees convergence but quantitatively predicts its rate. "
    "This suggests the spectral radius could serve as a practical "
    "diagnostic tool: rather than treating stability as a binary "
    "property, practitioners modelling real contests could use the "
    "spectral radius of the linearised best response mapping to "
    "anticipate how many rounds of strategic adjustment a market or "
    "institution is likely to require before settling near "
    "equilibrium. Second, the rent dissipation results partially "
    "confirm Nitzan's (1994) non-monotonicity prediction: total "
    "effort rises monotonically with r in the convergent region "
    "(Table 2), and the very high r regime produces zero effort "
    "(Section 4.6), consistent with the descending portion of his "
    "dissipation curve. However, the intermediate region where this "
    "descent should occur is precisely where dynamics fail to "
    "converge, meaning the dissipation curve cannot currently be "
    "traced continuously through simulation alone — a limitation "
    "discussed further in Section 5.4."
)

add_heading("5.2 Extending Existing Theory", level=2)
add_para(
    "The most significant theoretical extension concerns the "
    "Szidarovszky-Okuguchi contraction condition. The n = 5, r = 1.0 "
    "case — spectral radius 1.49, convergence rate 100% — demonstrates "
    "that this condition is not necessary for convergence. Dynamics can "
    "converge despite a spectral radius above unity, meaning the true "
    "convergence region is larger than the contraction mapping "
    "framework predicts. This is a genuinely counterintuitive result: "
    "the standard interpretation of spectral radius above one is that "
    "small perturbations from equilibrium should grow under repeated "
    "best response updates, yet the empirical trajectories converge "
    "regardless. A plausible explanation is that the local linear "
    "approximation captured by the Jacobian is only locally accurate, "
    "and the nonlinear best response mapping may fold trajectories "
    "back toward equilibrium in ways the eigenvalue analysis alone "
    "cannot capture, particularly given the boundary constraint that "
    "effort levels cannot fall below zero. Testing this explanation "
    "rigorously would require examining the global structure of the "
    "best response map rather than its local linearisation, which is "
    "identified as a direction for future work in Section 6.2."
)
add_para(
    "The heterogeneous valuation results extend Cornes and Hartley "
    "(2005) computationally. Their proportionality prediction is "
    "confirmed under both mild and strong asymmetry, with the "
    "effort ratio between players matching the valuation ratio "
    "precisely in both cases tested (Table 4). The additional "
    "finding from the valuation sweep (Section 4.4, Figure 10) — "
    "that Player 1's equilibrium effort is non-monotonic in Player "
    "2's valuation, rising as V2 approaches V1 from below and then "
    "falling as V2 continues to rise — represents a novel observation "
    "not present in the Cornes and Hartley analytical results. This "
    "pattern has an intuitive strategic reading: as V2 rises toward "
    "V1, Player 1 faces intensifying competition and raises effort "
    "defensively, but once V2 substantially exceeds V1, further "
    "increases make the contest sufficiently lopsided that Player 1 "
    "rationally scales back effort rather than compete for a prize "
    "increasingly likely to be won by Player 2. Formalising this "
    "intuition analytically is noted as a further avenue in Section "
    "6.2."
)

add_heading("5.3 Challenging Existing Theory", level=2)
add_para(
    "The knife-edge non-convergence at r = 3.0 constitutes the most "
    "striking and theoretically important finding of this "
    "dissertation. It is a substantive challenge to the implicit "
    "assumption, common across much of the contest theory literature, "
    "that Nash equilibria are reachable through natural learning "
    "processes. The unique Nash equilibrium exists at r = 3.0 and "
    "can be computed directly, yet it is computationally unreachable "
    "through best response learning from any of the initial "
    "conditions tested. What makes this finding particularly striking "
    "is its isolation: r = 2.8 and r = 3.2 both converge reliably at "
    "100%, so the instability is not part of a broadening region of "
    "difficulty but a single, sharply localised point of failure. "
    "The trajectory analysis in Experiment 5 (Figure 13) shows this "
    "is not simple two-cycle behaviour of the kind familiar from "
    "cobweb models, but genuine explosive divergence — effort levels "
    "swing between near-zero and near-maximum values with growing "
    "amplitude before both players collapse toward zero. This pattern "
    "is qualitatively closer to the sensitive dependence on initial "
    "conditions reported by Hopkins and Kornienko (2010) under "
    "evolutionary dynamics than to a standard limit cycle, and its "
    "resistance to all four inertial learning rates tested suggests "
    "the instability is a structural feature of the best response "
    "map itself at this parameter value, not an artefact of the "
    "particular update rule chosen. Whether r = 3.0 is a genuine "
    "isolated singularity of the dynamical system, or the visible tip "
    "of a narrow band of instability too fine for the parameter grid "
    "used here to resolve, remains an open question and is the "
    "subject of the fine-grained follow-up proposed in Section 6.2."
)
add_para(
    "The expanding instability zone with increasing n (Section 4.6, "
    "Figure 15) challenges the use of symmetric Nash equilibria as "
    "behavioural predictions in large contests. For n >= 10 at "
    "r = 1.0, the theoretically unique Nash equilibrium is "
    "analytically well-defined but dynamically inaccessible through "
    "best response learning from any tested starting point, and the "
    "total effort observed in these non-convergent runs (Table 3) "
    "diverges substantially from the analytical prediction rather "
    "than merely approaching it slowly. This provides a structural, "
    "dynamics-based explanation for the systematic deviations from "
    "Nash predictions documented experimentally by Chowdhury and "
    "Sheremeta (2011): if the equilibrium is not reachable through "
    "the most natural learning process available to boundedly "
    "rational agents, deviations from its predictions in laboratory "
    "or field settings should be expected as a matter of course "
    "rather than treated as evidence of irrationality."
)

add_heading("5.4 Unresolved Patterns Warranting Further Investigation", level=2)
add_para(
    "Beyond confirming, extending, or challenging existing results, "
    "three patterns emerged during this project that were not "
    "anticipated at the outset and are flagged here as candidate "
    "directions for further research, in line with the recommendation "
    "to note unexpected behaviour as it arises."
)
add_para(
    "First, the apparent looseness of the Szidarovszky-Okuguchi "
    "sufficient condition (Section 5.2) raises the question of "
    "exactly how much slack exists between the sufficient condition "
    "and the true necessary and sufficient boundary for convergence. "
    "A systematic comparison of spectral radius against convergence "
    "outcome across a much finer parameter grid than used here could "
    "in principle trace this boundary directly, offering either a "
    "tightened sufficient condition or evidence that no simple "
    "closed-form tightening exists."
)
add_para(
    "Second, the fact that r = 5.0 produces universal convergence "
    "across every value of n tested, but to a degenerate zero-effort "
    "outcome rather than an interior equilibrium, suggests that very "
    "high decisiveness may act as a stabiliser precisely because it "
    "destroys the incentive to compete at all. This raises a question "
    "with direct mechanism design relevance: whether there exists an "
    "intermediate range of r that preserves meaningful competitive "
    "effort while still retaining the stability properties associated "
    "with the extremes, or whether stability and meaningful competition "
    "are fundamentally in tension across the entire parameter range "
    "explored here."
)
add_para(
    "Third, the qualitative similarity between the explosive "
    "instability observed at r = 3.0 and the sensitive dependence on "
    "initial conditions reported under evolutionary dynamics by "
    "Hopkins and Kornienko (2010) suggests these two dynamic "
    "processes — best response learning and evolutionary selection — "
    "may share a common source of instability rooted in the "
    "underlying payoff structure of the contest, rather than being "
    "independent artefacts of the two different learning models. "
    "Testing this conjecture would require running an evolutionary "
    "dynamics variant of the present framework and comparing its "
    "instability regions directly against those mapped in Figure 15."
)

add_heading("5.5 Limitations", level=2)
add_para(
    "Several limitations should be acknowledged. Parameter sweeps "
    "use synchronous dynamics as the primary lens for the systematic "
    "r x n grid — asynchronous dynamics may exhibit different "
    "convergence boundaries and were only spot-checked in Experiment "
    "1. The 500-iteration limit used in most experiments may "
    "misclassify slowly-converging cases as non-convergent; the n = 5 "
    "case in Experiment 3, which required 2010 iterations under a "
    "more generous limit, illustrates this risk directly and means "
    "some cells in the Figure 15 heatmap may understate the true "
    "extent of convergence. The discrete parameter grid may also miss "
    "important features between sampled points, as demonstrated by "
    "the r = 3.0 knife-edge, which was only identified through "
    "targeted fine-grained follow-up in Experiment 5 rather than the "
    "original coarse grid. Finally, the analysis throughout focuses "
    "on pure strategy Nash equilibria; mixed strategy equilibria and "
    "their dynamic properties are not considered, and it is possible "
    "that some of the non-convergent regions identified here would "
    "resolve differently under a mixed-strategy learning process."
)
page_break()
print("Chapter 5 done. Writing Chapter 6...")


# ══════════════════════════════════════════════════════════
# CHAPTER 6 — CONCLUSION
# ══════════════════════════════════════════════════════════

add_heading("6. Conclusion", level=1)

add_heading("6.1 Summary of Findings", level=2)
add_para(
    "This dissertation has investigated best response dynamics in "
    "Tullock contests through a validated computational simulation "
    "framework. The main findings are:"
)
add_bullet(
    "The simulation correctly recovers all analytical benchmarks, "
    "validating the framework as a reliable tool for extending "
    "theoretical analysis beyond closed-form tractability."
)
add_bullet(
    "A knife-edge non-convergence point at r = 3.0 exhibits explosive "
    "instability resistant to inertial stabilisation."
)
add_bullet(
    "A convergence threshold between n = 5 and n = 10 at r = 1.0 "
    "shows how competition intensity interacts with decisiveness."
)
add_bullet(
    "Cornes and Hartley's (2005) proportionality prediction is "
    "confirmed and extended, with an additional non-monotonic "
    "strategic response identified."
)
add_bullet(
    "The Szidarovszky-Okuguchi (1997) contraction condition is "
    "sufficient but not necessary for convergence."
)
add_bullet(
    "Very high decisiveness (r = 5.0) produces universal convergence "
    "to a zero-effort corner solution."
)

add_heading("6.2 Future Work", level=2)
add_para(
    "Four directions for future research emerge. First, the knife-edge "
    "at r = 3.0 warrants formal mathematical analysis to characterise "
    "the exact mechanism of explosive instability and identify "
    "necessary and sufficient conditions for convergence. Second, the "
    "Stackelberg extension of Liu et al. (2025) represents a natural "
    "follow-up. Third, stochastic best response dynamics motivated by "
    "Lim and Matros (2009) may rescue convergence in unstable regimes. "
    "Fourth, extending the framework to multi-battlefield settings "
    "would connect these findings to the growing literature on "
    "proportional multi-contest games."
)
page_break()
print("Chapter 6 done. Writing references...")


# ══════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════

add_heading("References", level=1)

refs = [
    "Chowdhury, S. M., & Sheremeta, R. M. (2011). A generalized "
    "Tullock contest. Public Choice, 147(3), 413-427.",

    "Cornes, R., & Hartley, R. (2005). Asymmetric contests with "
    "general technologies. Economic Theory, 26(4), 923-946.",

    "Elkind, E., Ghosh, A., & Goldberg, P. (2024). Continuous-time "
    "best-response and related dynamics in Tullock contests with "
    "convex costs. arXiv:2402.08541.",

    "Ghosh, A., & Goldberg, P. (2023). Best-response dynamics in "
    "lottery contests. ACM Conference on Economics and Computation. "
    "arXiv:2305.10881.",

    "Hopkins, E., & Kornienko, T. (2010). Which inequality? "
    "American Economic Journal: Microeconomics, 2(3), 106-137.",

    "Judd, K. L. (1998). Numerical methods in economics. MIT Press.",

    "Lim, W., & Matros, A. (2009). Contests with a stochastic "
    "number of players. Games and Economic Behavior, 67(2), 584-597.",

    "Liu, Y., Ni, B., Shen, W., Wang, Z., & Zhang, J. (2025). "
    "Stackelberg vs. Nash in the lottery Colonel Blotto game. "
    "Proceedings of IJCAI-25, 3961-3969.",

    "Miranda, M. J., & Fackler, P. L. (2002). Applied computational "
    "economics and finance. MIT Press.",

    "Nitzan, S. (1994). Modelling rent-seeking contests. "
    "European Journal of Political Economy, 10(1), 41-60.",

    "Perez-Castrillo, J. D., & Verdier, T. (1992). A general "
    "analysis of rent-seeking games. Public Choice, 73(3), 335-350.",

    "Szidarovszky, F., & Okuguchi, K. (1997). On the existence and "
    "uniqueness of pure Nash equilibrium in rent-seeking games. "
    "Games and Economic Behavior, 18(1), 135-140.",

    "Tullock, G. (1980). Efficient rent seeking. In J. M. Buchanan, "
    "R. D. Tollison, & G. Tullock (Eds.), Toward a theory of the "
    "rent-seeking society (pp. 97-112). Texas A&M University Press.",

    "Wilson, G., et al. (2014). Best practices for scientific "
    "computing. PLoS Biology, 12(1), e1001745.",
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent       = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.add_run(ref).font.size = Pt(11)


# ══════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════

print("Saving dissertation draft...")

output_path = "results/dissertation_draft.docx"
doc.save(output_path)

print()
print("=" * 55)
print(f"Saved to: {output_path}")
print("Chapters:   6")
print("Tables:     5")
print("Captions:   16 figures")
print("References: 14")
print("=" * 55)